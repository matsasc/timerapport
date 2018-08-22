from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import date_utility as _date

#Word works with EMU units,
#EMU multiplier is what you need to multiply with cm to get the correct EMU value.
EMU_MULTIPLIER = 360000

def addRowToTable(table, rowTuple):
    row = table.add_row()
    for e in rowTuple:
        row.cells[e[0]].text = str(e[1])
        for p in row.cells[e[0]].paragraphs:
            adjustFont(p, 11, "Calibri")

def addColumnsToTable(table, widths):
    for width in widths:
        table.add_column(int(width*EMU_MULTIPLIER))


def adjustCellWidth(table, row, col, value):
    table.cell(row, col).width = int(value*EMU_MULTIPLIER)

def adjustTableWidth(table, widthTuple):
    for row in table.rows:
        for e in konsulentTabellWidth:
            row.cells[e[0]].width = e[1]*EMU_MULTIPLIER

def adjustFont(text, size, fontName):
    for run in text.runs:
        font = run.font
        font.size = Pt(size)
        font.name = fontName

document = Document()

title = document.add_paragraph('Timerapport')
adjustFont(title, 26, "Calibri")

periode = ["Periode:", _date.month_range()]
konsulent = ["Konsulent:", "TODO"]
oppdrag = ["Oppdrag:", "TODO"]
kunde = ["Kunde:", "TODO"]

konsulentRows = map(enumerate, [periode, konsulent, oppdrag, kunde])
konsulentTabellTuppel = list(enumerate(["Periode:", "Konsulent:", "Oppdrag:", "Kunde:"]))
konsulentTabellWidth = [3.05, 9.13]

konsulentTabell = document.add_table(0,0)
addColumnsToTable(konsulentTabell, konsulentTabellWidth)
for row in konsulentRows:
    addRowToTable(konsulentTabell, row)
konsulentTabell.style = "Table Grid"

document.add_paragraph('')
sammendrag = document.add_paragraph('Sammendrag')
adjustFont(sammendrag, 14, "Calibri")

uker = _date.week_range()

sammendragTabellWidth = [3.19, 5, 4]
sammendragBody = map(enumerate, uker)
sammendragHead = list(enumerate(["Uke",	"Dato", "Timer"]))
sammendragFoot = list(enumerate(["TOTALT", _date.last_date(), 0]))
sammendragRows = list([sammendragHead, *sammendragBody, sammendragFoot])

sammendragTabell = document.add_table(0,0)
addColumnsToTable(sammendragTabell, sammendragTabellWidth)
for row in sammendragRows:
    addRowToTable(sammendragTabell, row)
sammendragTabell.style = "Table Grid"

document.add_paragraph('')

for uke in uker:
    p = document.add_paragraph("Uke {0} - {1}".format(uke[0], uke[1]))
    adjustFont(p, 14, "Calibri")
    
    document.add_picture('placeholder.gif', width=Inches(3))
    document.add_paragraph('')

document.save('generated.docx')